"""
FinBot — backend/routes/upload.py
==================================
Document upload endpoint — parses uploaded files, chunks them,
embeds into a session-scoped Chroma collection for RAG retrieval.

Endpoints:
    POST /upload  (multipart: file + session_id)
"""

import os
import logging
from pathlib import Path

from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from backend.schemas import UploadResponse
from backend.db import db

logger = logging.getLogger(__name__)
router = APIRouter()

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv"}
CHROMA_DB_PATH = os.getenv("CHROMA_DB_PATH", "./chroma_db")
EMBEDDING_MODEL = "BAAI/bge-base-en-v1.5"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


def _extract_text(file_path: str, ext: str) -> tuple:
    """Extract text from a file. Returns (text, metadata_dict)."""
    if ext == ".pdf":
        import fitz
        doc = fitz.open(file_path)
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                pages.append((text, {"page_number": i + 1, "total_pages": len(doc)}))
        doc.close()
        if not pages:
            return "", {"page_count": 0}
        full_text = "\n\n".join(t for t, _ in pages)
        return full_text, {"page_count": len(pages)}

    elif ext == ".docx":
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        return full_text, {"paragraph_count": len(paragraphs)}

    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            full_text = f.read()
        return full_text, {"char_count": len(full_text)}

    elif ext == ".csv":
        import pandas as pd
        df = pd.read_csv(file_path)
        full_text = df.to_string(index=False)
        return full_text, {"row_count": len(df), "col_count": len(df.columns)}

    raise ValueError(f"Unsupported file type: {ext}")


def _chunk_text(text: str) -> list:
    """Split text into chunks using RecursiveCharacterTextSplitter."""
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", "! ", "? ", ", ", " ", ""],
    )
    doc = Document(page_content=text)
    return splitter.split_documents([doc])


def _embed_to_session_collection(session_id: str, chunks: list) -> None:
    """Embed chunks into a session-scoped Chroma collection."""
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_chroma import Chroma

    collection_name = f"session_{session_id}"

    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=collection_name,
        persist_directory=CHROMA_DB_PATH,
    )
    logger.info("Embedded %d chunks into collection '%s'", len(chunks), collection_name)


@router.post("/upload", response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    session_id: str = Form(...),
) -> UploadResponse:
    """
    Upload a document to augment a specific session's retrieval.

    The file is parsed, chunked, and embedded into a session-scoped
    Chroma collection named ``session_{session_id}``. Subsequent chat
    requests in this session retrieve from both the global knowledge
    base and this session collection.
    """
    # Validate session exists (create if not)
    if not db.get_session(session_id):
        db.create_session(session_id, title="")

    # Validate file extension
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Read and validate size
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large ({len(content) // (1024*1024)}MB). Max: 20MB",
        )

    # Save to temp location
    import tempfile
    tmp_dir = tempfile.mkdtemp()
    tmp_path = os.path.join(tmp_dir, file.filename)
    with open(tmp_path, "wb") as f:
        f.write(content)

    try:
        # Extract text
        text, meta = _extract_text(tmp_path, ext)
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the file")

        # Chunk
        chunks = _chunk_text(text)
        if not chunks:
            raise HTTPException(status_code=400, detail="Document produced no chunks after splitting")

        # Add metadata to each chunk
        from langchain_core.documents import Document
        for i, chunk in enumerate(chunks):
            chunk.metadata["file_name"] = file.filename
            chunk.metadata["source"] = "uploaded"
            chunk.metadata["session_id"] = session_id

        # Embed into session collection
        _embed_to_session_collection(session_id, chunks)

        # Ensure session has a title derived from filename
        session = db.get_session(session_id)
        if session and not session["title"]:
            db.update_session_title(session_id, file.filename[:100])

        page_count = meta.get("page_count", meta.get("paragraph_count", meta.get("row_count", 0)))
        logger.info(
            "Uploaded '%s' — %d chunks, %s pages/rows",
            file.filename, len(chunks), page_count,
        )

        return UploadResponse(
            filename=file.filename,
            chunk_count=len(chunks),
            page_count=page_count,
            status="success",
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error("Upload failed: %s", str(e))
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")
    finally:
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)
